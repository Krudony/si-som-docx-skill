import os
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

def create_perfect_template(path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Twips(1440)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH SarabunPSK'
    font.size = Pt(16)

    def add_line(text, bold=False, align=None):
        p = doc.add_paragraph()
        if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        return p

    add_line('โครงการ                     {{ project_name }}', bold=True)
    add_line('แผนงาน                      {{ plan_name }}', bold=True)
    add_line('สนองกลยุทธ์สถานศึกษา           {{ strategy }}', bold=True)
    add_line('มาตรฐานดำเนินการ             {{ standard }}', bold=True)
    add_line('ลักษณะโครงการ                {{ type }}', bold=True)
    add_line('งบประมาณ                         {{ budget_total }}  บาท', bold=True)
    add_line('ผู้รับผิดชอบโครงการ          {{ responsible }}', bold=True)
    add_line('ระยะเวลาดำเนินการ            {{ duration }}', bold=True)
    add_line('********************************************', align='center')

    add_line('1.  หลักการและเหตุผล', bold=True)
    add_line('             {{ principle_text }}')
    add_line('2.  วัตถุประสงค์', bold=True)
    add_line('{{ objectives_text }}')
    add_line('3.   เป้าหมาย', bold=True)
    add_line('     3.1  เชิงปริมาณ', bold=True)
    add_line('{{ target_quant }}')
    add_line('     3.2  เชิงคุณภาพ', bold=True)
    add_line('{{ target_qual }}')

    add_line('4. กิจกรรมและขั้นตอนการดำเนินโครงการ', bold=True)
    t0 = doc.add_table(rows=2, cols=4)
    t0.style = 'Table Grid'
    row0 = t0.rows[1].cells
    row0[0].text = '{% for a in activities %}{{ a.step }}'; row0[1].text = '{{ a.time }}'; row0[2].text = '{{ a.owner }}'; row0[3].text = '{{ a.eval }}{% endfor %}'

    add_line('\n5.  งบประมาณ        {{ budget_source }}', bold=True)
    t1 = doc.add_table(rows=2, cols=6)
    t1.style = 'Table Grid'
    row1 = t1.rows[1].cells
    row1[0].text = '{% for b in budget_items %}{{ b.id }}'; row1[1].text = '{{ b.name }}'; row1[2].text = '{{ b.amount }}'; row1[3].text = '{{ b.reward }}'; row1[4].text = '{{ b.use }}'; row1[5].text = '{{ b.mat }}{% endfor %}'
    
    add_line('\n6.  การประเมินผล', bold=True)
    t2 = doc.add_table(rows=2, cols=4)
    t2.style = 'Table Grid'
    row2 = t2.rows[1].cells
    row2[0].text = '{% for e in evals %}{{ e.id }}'; row2[1].text = '{{ e.index }}'; row2[2].text = '{{ e.method }}'; row2[3].text = '{{ e.tool }}{% endfor %}'

    add_line('\n7.    ผลที่คาดว่าจะได้รับ', bold=True)
    add_line('{{ expected_results }}')

    add_line('\n\n')
    add_line('                     ผู้เสนอโครงการ                                                ผู้เห็นชอบโครงการ      ', align='center')
    add_line('\n')
    add_line('                ( นายอุดร  มะโนเนือง)                                        (นายสมศักดิ์  อิปิน)       ', align='center')
    add_line('                 ตำแหน่ง ครูชำนาญการ                         ประธานคณะกรรมการสถานศึกษาขั้นพื้นฐาน', align='center')
    add_line('\n')
    add_line('           ผู้อนุมัติโครงการ', bold=True, align='center')
    add_line('\n')
    add_line('                                               (นายสงวน  จันทอน)', align='center')
    add_line('                             ผู้อำนวยการโรงเรียนบ้านบ้านแม่ทราย(คุรุราษฎร์เจริญวิทย์)', align='center')
    doc.save(path)

def render_final(template_path, output_path):
    doc = DocxTemplate(template_path)
    context = {
        'project_name': 'การพัฒนาอาคารสถานที่และสิ่งแวดล้อมอย่างมีคุณภาพ',
        'plan_name': 'งานบริหารทั่วไป',
        'strategy': 'กลยุทธ์ที่ 4',
        'standard': 'มฐ.ที่ 2.5',
        'type': 'โครงการต่อเนื่อง',
        'budget_total': '10,000',
        'responsible': 'นายอุดร มะโนเนือง',
        'duration': 'ปีการศึกษา 2569',
        'principle_text': 'อาคารสถานที่และสิ่งแวดล้อมในสถานศึกษา เป็นปัจจัยสำคัญที่ส่งผลต่อการจัดการเรียนรู้อย่างมีคุณภาพ สถานศึกษาที่สะอาด ร่มรื่น ปลอดภัย และมีอาคารสถานที่ที่มั่นคงแข็งแรง จะช่วยสร้างบรรยากาศที่เอื้อต่อการเรียนรู้ของนักเรียนและส่งเสริมสุขภาพจิตที่ดีของคณะครูและบุคลากรทางการศึกษา โรงเรียนจึงต้องมีการบริหารจัดการอาคารสถานที่ให้พร้อมใช้งานตลอดเวลา',
        'objectives_text': '            2.1  เพื่อปรับปรุงและพัฒนาอาคารสถานที่และสิ่งแวดล้อมให้มีความมั่นคง แข็งแรง และปลอดภัย\n            2.2  เพื่อสร้างบรรยากาศและสิ่งแวดล้อมที่เอื้อต่อการจัดการเรียนรู้ของนักเรียน',
        'target_quant': '                   3.1.1  อาคารสถานที่และห้องเรียนได้รับการปรับปรุงและซ่อมแซมให้พร้อมใช้งาน ร้อยละ 95',
        'target_qual': '                    3.2.1 โรงเรียนมีสภาพแวดล้อมที่เอื้อต่อการเรียนรู้ มีความปลอดภัย',
        'activities': [{'step': '1. วางแผน', 'time': 'ก.ค. 69', 'owner': 'อุดร', 'eval': 'บันทึก'}, {'step': '2. ทำงาน', 'time': 'ส.ค. 69', 'owner': 'อุดร', 'eval': 'รูป'}],
        'budget_source': 'งบอุดหนุน',
        'budget_items': [{'id': '1', 'name': 'วัสดุ', 'amount': '10,000', 'reward': '-', 'use': '-', 'mat': '10,000'}],
        'evals': [{'id': '1', 'index': 'ความปลอดภัย', 'method': 'สังเกต', 'tool': 'แบบประเมิน'}],
        'expected_results': '              7.1 สภาพแวดล้อมสวยงาม ปลอดภัย'
    }
    doc.render(context)
    doc.save(output_path)

if __name__ == '__main__':
    t_path = r'C:\Users\User\Desktop\New folder (3)\PERFECT_TEMPLATE.docx'
    o_path = r'C:\Users\User\Desktop\New folder (3)\สรุปโครงการ_อาคารสถานที่_สมบูรณ์แบบ.docx'
    create_perfect_template(t_path)
    render_final(t_path, o_path)
