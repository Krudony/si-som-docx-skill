import sys
import os
from docx import Document
from docx.shared import Pt, Twips
from docxtpl import DocxTemplate

def create_template(path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Twips(1418)

    def add_para(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(16)
        run.bold = bold
        return p

    add_para('โครงการ: {{ project_name }}', bold=True)
    add_para('ผู้รับผิดชอบ: {{ responsible }}')
    add_para('งบประมาณ: {{ budget }} บาท')
    add_para('--- ตารางกิจกรรม ---', bold=True)
    
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'กิจกรรม'
    hdr_cells[1].text = 'เวลา'
    hdr_cells[2].text = 'คนรับผิดชอบ'
    
    # Simple Loop inside Table
    row_cells = table.rows[1].cells
    row_cells[0].text = '{% for item in activities %}{{ item.name }}'
    row_cells[1].text = '{{ item.time }}'
    row_cells[2].text = '{{ item.owner }}{% endfor %}'

    doc.save(path)
    print(f'Template created: {path}')

def render_project(template_path, output_path):
    doc = DocxTemplate(template_path)
    context = {
        'project_name': 'การพัฒนาอาคารสถานที่ (DocxTpl Demo)',
        'responsible': 'นายอุดร มะโนเนือง',
        'budget': '10,000',
        'activities': [
            {'name': '1. บิ๊กคลีนนิ่ง', 'time': 'ก.ค. 69', 'owner': 'อุดร'},
            {'name': '2. ปรับปรุงสวน', 'time': 'ส.ค. 69', 'owner': 'ทีมงาน'}
        ]
    }
    doc.render(context)
    doc.save(output_path)
    print(f'Final Document Rendered: {output_path}')

if __name__ == '__main__':
    t_path = r'C:\Users\User\Desktop\New folder (3)\template_test.docx'
    o_path = r'C:\Users\User\Desktop\New folder (3)\สรุปโครงการ_DocxTpl_Final.docx'
    create_template(t_path)
    render_project(t_path, o_path)
