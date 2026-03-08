import sys
import os
from docx import Document
from docx.shared import Twips, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION     
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class SiSomDocxEngine:
    def __init__(self, filename=None, font_size=16):
        if filename and os.path.exists(filename): self.doc = Document(filename)
        else: self.doc = Document(); self._set_page_setup(); self.set_font('TH SarabunPSK', font_size)

    def _set_page_setup(self):
        section = self.doc.sections[0]
        section.page_height = Twips(16838); section.page_width = Twips(11906)
        section.top_margin = Twips(1418); section.bottom_margin = Twips(1418)
        section.left_margin = Twips(1418); section.right_margin = Twips(1418)

    def set_font(self, name='TH SarabunPSK', size=16):
        style = self.doc.styles['Normal']
        font = style.font; font.name = name; font.size = Pt(size)
        rPr = style._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), name); rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:eastAsia'), name); rFonts.set(qn('w:cs'), name)
        rPr.append(rFonts)

    def add_paragraph(self, text, bold=False, align=None, size=None):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        if bold: run.bold = True
        if size: run.font.size = Pt(size)
        if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return p

    def add_table(self, rows, cols, data):
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = cell_text
                # Set font for table cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'TH SarabunPSK'
                        run.font.size = Pt(16)
        return table

    def save(self, filename):
        self.doc.save(filename)
        print(f'Document saved: {filename}')

if __name__ == '__main__':
    engine = SiSomDocxEngine()
    
    # Header Info
    engine.add_paragraph('โครงการ                 การพัฒนาอาคารสถานที่และสิ่งแวดล้อมอย่างมีคุณภาพ', bold=True)
    engine.add_paragraph('แผนงาน                  งานบริหารทั่วไป', bold=True)
    engine.add_paragraph('สนองกลยุทธ์สถานศึกษา           กลยุทธ์ที่ 4', bold=True)
    engine.add_paragraph('มาตรฐานดำเนินการ                มฐ.ที่ 2.5', bold=True)
    engine.add_paragraph('ลักษณะโครงการ           โครงการต่อเนื่อง', bold=True)
    engine.add_paragraph('งบประมาณ                         10,000  บาท', bold=True)
    engine.add_paragraph('ผู้รับผิดชอบโครงการ             นายอุดร  มะโนเนือง', bold=True)
    engine.add_paragraph('ระยะเวลาดำเนินการ               ปีการศึกษา 2569', bold=True)
    engine.add_paragraph('********************************************', align='center')

    # 1. Principles
    engine.add_paragraph('1.  หลักการและเหตุผล', bold=True)
    engine.add_paragraph('             อาคารสถานที่และสิ่งแวดล้อมในสถานศึกษา เป็นปัจจัยสำคัญที่ส่งผลต่อการจัดการเรียนรู้อย่างมีคุณภาพ สถานศึกษาที่สะอาด ร่มรื่น ปลอดภัย และมีอาคารสถานที่ที่มั่นคงแข็งแรง จะช่วยสร้างบรรยากาศที่เอื้อต่อการเรียนรู้ของนักเรียนและส่งเสริมสุขภาพจิตที่ดีของคณะครูและบุคลากรทางการศึกษา โรงเรียนจึงต้องมีการบริหารจัดการอาคารสถานที่ให้พร้อมใช้งานตลอดเวลา')

    # 2. Objectives
    engine.add_paragraph('2.  วัตถุประสงค์', bold=True)
    engine.add_paragraph('            2.1  เพื่อปรับปรุงและพัฒนาอาคารสถานที่และสิ่งแวดล้อมให้มีความมั่นคง แข็งแรง และปลอดภัย')
    engine.add_paragraph('            2.2  เพื่อสร้างบรรยากาศและสิ่งแวดล้อมที่เอื้อต่อการจัดการเรียนรู้ของนักเรียน')

    # 3. Goals
    engine.add_paragraph('3.   เป้าหมาย', bold=True)
    engine.add_paragraph('         3.1  เชิงปริมาณ', bold=True)
    engine.add_paragraph('                   3.1.1  อาคารสถานที่และห้องเรียนได้รับการปรับปรุงและซ่อมแซมให้พร้อมใช้งาน ร้อยละ 95')
    engine.add_paragraph('         3.2  เชิงคุณภาพ', bold=True)
    engine.add_paragraph('                    3.2.1 โรงเรียนมีสภาพแวดล้อมที่เอื้อต่อการเรียนรู้ มีความปลอดภัย')

    # 4. Activities Table
    engine.add_paragraph('4. กิจกรรมและขั้นตอนการดำเนินโครงการ', bold=True)
    act_data = [
        ['ขั้นตอนการดำเนินงาน', 'ระยะเวลา', 'ผู้รับผิดชอบ', 'การประเมินผล'],
        ['1. ขั้นวางแผน(P)\n1.1 ประชุมวางแผน\n1.2 แต่งตั้งคณะทำงาน\n2. ขั้นดำเนินการ(D)\n2.1 ปรับปรุงภูมิทัศน์\n2.2 ซ่อมแซมอาคารเรียน\n3. ขั้นประเมินผล(C)\n3.1 สรุปผลโครงการ\n4. ขั้นปรับปรุง(A)\n4.1 วิเคราะห์ผลเพื่อพัฒนาต่อ', 'ตลอดปีการศึกษา 2569', 'นายอุดร มะโนเนือง', 'สังเกต/แบบสอบถาม']
    ]
    engine.add_table(2, 4, act_data)

    # 5. Budget Table
    engine.add_paragraph('\n5.  งบประมาณ', bold=True)
    bud_data = [
        ['ที่', 'รายการ', 'งบประมาณ', 'ตอบแทน', 'ใช้สอย', 'วัสดุ'],
        ['1', 'วัสดุปรับปรุงอาคาร', '10,000', '-', '-', '10,000'],
        ['รวม', 'รวม', '10,000', '-', '-', '10,000']
    ]
    engine.add_table(3, 6, bud_data)

    # 6. Evaluation Table
    engine.add_paragraph('\n6.  การประเมินผล', bold=True)
    eval_data = [
        ['ที่', 'ดัชนีบ่งชี้ความสำเร็จ', 'วิธีการ/ประเมินผล', 'เครื่องมือที่ใช้วัด/ประเมินผล'],
        ['1', 'อาคารสถานที่ปลอดภัย\nร้อยละ 95', 'สังเกต/ตรวจสอบ', 'แบบประเมินความพึงพอใจ']
    ]
    engine.add_table(2, 4, eval_data)

    # 7. Expected Results
    engine.add_paragraph('\n7.    ผลที่คาดว่าจะได้รับ', bold=True)
    engine.add_paragraph('              7.1 โรงเรียนมีสภาพแวดล้อมที่สวยงาม ปลอดภัย เอื้อต่อการเรียนรู้')

    # Signatures
    engine.add_paragraph('\n\n')
    engine.add_paragraph('                     ผู้เสนอโครงการ                                                ผู้เห็นชอบโครงการ      ', align='center')
    engine.add_paragraph('\n')
    engine.add_paragraph('                ( นายอุดร  มะโนเนือง)                                        (นายสมศักดิ์  อิปิน)       ', align='center')
    engine.add_paragraph('                 ตำแหน่ง ครูชำนาญการ                         ประธานคณะกรรมการสถานศึกษาขั้นพื้นฐาน', align='center')
    engine.add_paragraph('\n')
    engine.add_paragraph('           ผู้อนุมัติโครงการ', align='center', bold=True)
    engine.add_paragraph('\n')
    engine.add_paragraph('                                               (นายสงวน  จันทอน)', align='center')
    engine.add_paragraph('                             ผู้อำนวยการโรงเรียนบ้านบ้านแม่ทราย(คุรุราษฎร์เจริญวิทย์)', align='center')

    engine.save(r'C:\Users\User\Desktop\New folder (3)\โครงการพัฒนาอาคารสถานที่_นายอุดร_v2.docx')
