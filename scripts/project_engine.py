import os
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class SchoolProjectEngine:
    def __init__(self, output_path):
        self.doc = Document()
        self.output_path = output_path
        self._set_page_setup()
        self._set_global_font()

    def _set_page_setup(self):
        section = self.doc.sections[0]
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Twips(1440)

    def _set_global_font(self):
        style = self.doc.styles['Normal']
        font = style.font; font.name = 'TH SarabunPSK'; font.size = Pt(16)
        pf = style.paragraph_format; pf.line_spacing = 1.0; pf.space_after = pf.space_before = Pt(0); pf.widow_control = False
        rPr = style._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'TH SarabunPSK'); rFonts.set(qn('w:hAnsi'), 'TH SarabunPSK'); rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK'); rFonts.set(qn('w:cs'), 'TH SarabunPSK')
        rPr.append(rFonts)

    def add_line(self, text, bold=False, align=None, indent=False):
        p = self.doc.add_paragraph()
        if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'left': p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else: p.alignment = getattr(WD_ALIGN_PARAGRAPH, 'THAI_JUSTIFY', WD_ALIGN_PARAGRAPH.JUSTIFY)
        if indent: p.paragraph_format.first_line_indent = Twips(720)
        run = p.add_run(text); run.bold = bold
        return p

    def save(self):
        self.doc.save(self.output_path)
        print(f'Document saved: {self.output_path}')

def generate_full_project(path, ctx):
    engine = SchoolProjectEngine(path)
    engine.add_line(f"โครงการ                     {ctx['project_name']}", bold=True, align='left')
    engine.add_line(f"แผนงาน                      {ctx['plan_name']}", bold=True, align='left')
    engine.add_line(f"ผู้รับผิดชอบโครงการ          {ctx['responsible']}", bold=True, align='left')
    engine.add_line("********************************************", align='center')
    engine.add_line("1.  หลักการและเหตุผล", bold=True, align='left')
    for p in ctx['principles']: engine.add_line(p, indent=True)
    
    # Tables with correct logic
    engine.add_line("4. กิจกรรมและขั้นตอนการดำเนินโครงการ", bold=True, align='left')
    t0 = engine.doc.add_table(rows=2, cols=4); t0.style = 'Table Grid'
    hdr0 = t0.rows[0].cells; hdr0[0].text = 'ขั้นตอนการดำเนินงาน'; hdr0[1].text = 'ระยะเวลา'; hdr0[2].text = 'ผู้รับผิดชอบ'; hdr0[3].text = 'การประเมินผล'
    for cell in hdr0: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; cell.paragraphs[0].runs[0].bold = True
    r0 = t0.rows[1].cells; r0[0].text = ctx['act_steps']; r0[1].text = ctx['act_time']; r0[2].text = ctx['act_owner']; r0[3].text = ctx['act_eval']

    engine.add_line(f"\n5.  งบประมาณ        {ctx['budget_source']}", bold=True, align='left')
    t1 = engine.doc.add_table(rows=4, cols=6); t1.style = 'Table Grid'
    t1.cell(0, 0).text = 'ที่'; t1.cell(1, 0).text = 'ที่'; t1.cell(0, 1).text = 'รายการ'; t1.cell(1, 1).text = 'รายการ'; t1.cell(0, 2).text = 'งบประมาณ'; t1.cell(1, 2).text = 'งบประมาณ'; t1.cell(0, 3).text = 'จำแนกตามหมวดรายจ่าย'; t1.cell(1, 3).text = 'ตอบแทน'; t1.cell(1, 4).text = 'ใช้สอย'; t1.cell(1, 5).text = 'วัสดุ'
    t1.cell(0, 0).merge(t1.cell(1, 0)); t1.cell(0, 1).merge(t1.cell(1, 1)); t1.cell(0, 2).merge(t1.cell(1, 2)); t1.cell(0, 3).merge(t1.cell(0, 5))
    t1.cell(2, 0).text = '1'; t1.cell(2, 1).text = ctx['budget_item_name']; t1.cell(2, 2).text = ctx['budget_total']; t1.cell(2, 3).text = '-'; t1.cell(2, 4).text = '-'; t1.cell(2, 5).text = ctx['budget_total']
    t1.cell(3, 0).text = 'รวม'; t1.cell(3, 1).text = 'รวม'; t1.cell(3, 2).text = ctx['budget_total']; t1.cell(3, 3).text = '-'; t1.cell(3, 4).text = '-'; t1.cell(3, 5).text = ctx['budget_total']
    for r in range(2):
        for c in range(6):
            if t1.cell(r, c).paragraphs: t1.cell(r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; t1.cell(r, c).paragraphs[0].runs[0].bold = True

    # Global Font Sync
    for table in engine.doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0; paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs: run.font.name = 'TH SarabunPSK'; run.font.size = Pt(16)
    engine.save()

if __name__ == '__main__':
    sample_ctx = {'project_name': 'Test', 'plan_name': 'Test', 'responsible': 'Test', 'principles': ['Test'], 'act_steps': 'Test', 'act_time': 'Test', 'act_owner': 'Test', 'act_eval': 'Test', 'budget_source': 'Test', 'budget_item_name': 'Test', 'budget_total': '10,000'}
    generate_full_project(r'C:\Users\User\Desktop\New folder (3)\Engine_Test_Final.docx', sample_ctx)
