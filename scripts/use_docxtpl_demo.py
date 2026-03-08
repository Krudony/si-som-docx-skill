import sys
import os
from docx import Document
from docx.shared import Pt, Twips
from docxtpl import DocxTemplate

# 1. Create the Template File (Simulating a manually created Word template)
def create_template(path):
    doc = Document()
    
    # Set Margins
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Twips(1418)

    # Style helper
    def add_para(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(16)
        run.bold = bold
        return p

    add_para('โครงการ                 {{ project_name }}', bold=True)
    add_para('แผนงาน                  {{ plan_name }}')
    add_para('ผู้รับผิดชอบโครงการ             {{ responsible }}')
    add_para('งบประมาณ                         {{ budget }}  บาท')
    add_para('********************************************')
    
    add_para('1. กิจกรรมและขั้นตอนการดำเนินโครงการ', bold=True)
    
    # Create Table with Tags
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'กิจกรรม'
    hdr_cells[1].text = 'ระยะเวลา'
    hdr_cells[2].text = 'ผู้รับผิดชอบ'
    
    # Row with Loop Tags
    # Note: docxtpl uses special syntax inside table cells for loops
    row_cells = table.rows[1].cells
    row_cells[0].text = '{% tr for item in activities %}{{ item.name }}'
    row_cells[1].text = '{{ item.time }}'
    row_cells[2].text = '{{ item.owner }}{% tr endfor %}'

    doc.save(path)
    print(f'Template created: {path}')

# 2. Render the Template using docxtpl
def render_project(template_path, output_path):
    doc = DocxTemplate(template_path)
    
    context = {
        'project_name': 'การพัฒนาอาคารสถานที่และสิ่งแวดล้อมอย่างมีคุณภาพ',
        'plan_name': 'งานบริหารทั่วไป',
        'responsible': 'นายอุดร มะโนเนือง',
        'budget': '10,000',
        'activities': [
            {'name': '1. บิ๊กคลีนนิ่งเดย์', 'time': 'ก.ค. 69', 'owner': 'นายอุดร'},
            {'name': '2. ซ่อมแซมห้องน้ำ', 'time': 'ส.ค. 69', 'owner': 'นายอุดร'},
            {'name': '3. ปรับปรุงสวนหย่อม', 'time': 'ก.ย. 69', 'owner': 'ทีมงานอาคาร'}
        ]
    }
    
    doc.render(context)
    doc.save(output_path)
    print(f'Final Document Rendered: {output_path}')

if __name__ == '__main__':
    t_path = r'C:\Users\User\Desktop\New folder (3)\template_building.docx'
    o_path = r'C:\Users\User\Desktop\New folder (3)\สรุปโครงการ_ด้วย_DocxTpl.docx'
    create_template(t_path)
    render_project(t_path, o_path)
