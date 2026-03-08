import sys
import os
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, top=0, start=100, bottom=0, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{m[0]}')
        node.set(qn('w:w'), str(m[1]))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_manual_doc(output_path):
    doc = Document()
    
    # 1. Page Setup (914400 EMUs = 1440 Twips = 1 inch)
    section = doc.sections[0]
    section.top_margin = Twips(1440)
    section.bottom_margin = Twips(1440)
    section.left_margin = Twips(1440)
    section.right_margin = Twips(1440)

    # 2. Global Font Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH SarabunPSK'
    font.size = Pt(16)
    
    # Fix Thai Font in XML
    rPr = style._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'TH SarabunPSK')
    rFonts.set(qn('w:hAnsi'), 'TH SarabunPSK')
    rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    rFonts.set(qn('w:cs'), 'TH SarabunPSK')
    rPr.append(rFonts)

    def add_p(text, bold=False, align=None):
        p = doc.add_paragraph()
        if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        return p

    # 3. Content
    add_p('โครงการ                 การพัฒนาอาคารสถานที่และสิ่งแวดล้อมอย่างมีคุณภาพ', bold=True)
    add_p('แผนงาน                  งานบริหารทั่วไป', bold=True)
    add_p('สนองกลยุทธ์สถานศึกษา           กลยุทธ์ที่ 4', bold=True)
    add_p('มาตรฐานดำเนินการ             มฐ.ที่ 2.5', bold=True)
    add_p('ลักษณะโครงการ                โครงการต่อเนื่อง', bold=True)
    add_p('งบประมาณ                         10,000  บาท', bold=True)
    add_p('ผู้รับผิดชอบโครงการ          นายอุดร  มะโนเนือง', bold=True)
    add_p('ระยะเวลาดำเนินการ            ปีการศึกษา 2569', bold=True)
    add_p('********************************************', align='center')

    add_p('1.  หลักการและเหตุผล', bold=True)
    add_p('             อาคารสถานที่และสิ่งแวดล้อมในสถานศึกษา เป็นปัจจัยสำคัญที่ส่งผลต่อการจัดการเรียนรู้อย่างมีคุณภาพ สถานศึกษาที่สะอาด ร่มรื่น ปลอดภัย และมีอาคารสถานที่ที่มั่นคงแข็งแรง จะช่วยสร้างบรรยากาศที่เอื้อต่อการเรียนรู้ของนักเรียนและส่งเสริมสุขภาพจิตที่ดีของคณะครูและบุคลากรทางการศึกษา โรงเรียนจึงต้องมีการบริหารจัดการอาคารสถานที่ให้พร้อมใช้งานตลอดเวลา เพื่อยกระดับคุณภาพชีวิตและคุณภาพการศึกษาให้ยั่งยืน')

    add_p('2.  วัตถุประสงค์', bold=True)
    add_p('            2.1  เพื่อปรับปรุงและพัฒนาอาคารสถานที่และสิ่งแวดล้อมให้มีความมั่นคง แข็งแรง และปลอดภัย')
    add_p('            2.2  เพื่อสร้างบรรยากาศและสิ่งแวดล้อมที่เอื้อต่อการจัดการเรียนรู้ของนักเรียน')
    add_p('            2.3  เพื่อให้สถานศึกษามีความพร้อมในการให้บริการแก่ชุมชนและหน่วยงานภายนอก')

    add_p('3.   เป้าหมาย', bold=True)
    add_p('     3.1  เชิงปริมาณ', bold=True)
    add_p('                   3.1.1  อาคารสถานที่และห้องเรียนได้รับการปรับปรุงและซ่อมแซมให้พร้อมใช้งาน ร้อยละ 95')
    add_p('                   3.1.2  สภาพแวดล้อมโดยรอบมีความสะอาด ร่มรื่น และสวยงาม ร้อยละ 95')
    add_p('     3.2  เชิงคุณภาพ', bold=True)
    add_p('                    3.2.1 โรงเรียนมีสภาพแวดล้อมที่เอื้อต่อการเรียนรู้ มีความปลอดภัย และเป็นที่พึงพอใจของผู้รับบริการ')

    # 4. Table 0: Activities
    add_p('4. กิจกรรมและขั้นตอนการดำเนินโครงการ', bold=True)
    t0 = doc.add_table(rows=2, cols=4)
    t0.style = 'Table Grid'
    headers0 = ['ขั้นตอนการดำเนินงาน', 'ระยะเวลา', 'ผู้รับผิดชอบ', 'การประเมินผล']
    for i, h in enumerate(headers0):
        t0.cell(0, i).text = h
        t0.cell(0, i).paragraphs[0].runs[0].bold = True
        t0.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    content0 = [
        '1. ขั้นวางแผน(P)\n1.1 ประชุมวางแผน\n1.2 แต่งตั้งคณะทำงาน\n2. ขั้นดำเนินการ(D)\n2.1 ปรับปรุงภูมิทัศน์\n2.2 ซ่อมแซมอาคาร\n3. ขั้นประเมินผล(C)\n3.1 สรุปโครงการ\n4. ขั้นปรับปรุง(A)\n4.1 วิเคราะห์ผลเพื่อพัฒนาต่อ',
        'ตลอดปีการศึกษา 2569',
        'นายอุดร มะโนเนือง',
        'วิธีการประเมินผล\n1. สังเกต/ตรวจสอบ\nเครื่องมือ\n1. แบบประเมิน'
    ]
    for i, c in enumerate(content0):
        t0.cell(1, i).text = c

    # 5. Table 1: Budget (Complex Headers)
    add_p('\n5.  งบประมาณ        งบอุดหนุนรายหัว 10,000 บาท', bold=True)
    t1 = doc.add_table(rows=4, cols=6)
    t1.style = 'Table Grid'
    
    # Row 0 & 1 setup
    t1.cell(0, 0).text = 'ที่'; t1.cell(1, 0).text = 'ที่'
    t1.cell(0, 1).text = 'รายการ'; t1.cell(1, 1).text = 'รายการ'
    t1.cell(0, 2).text = 'งบประมาณ'; t1.cell(1, 2).text = 'งบประมาณ'
    t1.cell(0, 3).text = 'จำแนกตามหมวดรายจ่าย'; t1.cell(0, 4).text = 'จำแนกตามหมวดรายจ่าย'; t1.cell(0, 5).text = 'จำแนกตามหมวดรายจ่าย'
    t1.cell(1, 3).text = 'ตอบแทน'; t1.cell(1, 4).text = 'ใช้สอย'; t1.cell(1, 5).text = 'วัสดุ'
    
    # Merge cells for headers
    t1.cell(0, 0).merge(t1.cell(1, 0))
    t1.cell(0, 1).merge(t1.cell(1, 1))
    t1.cell(0, 2).merge(t1.cell(1, 2))
    t1.cell(0, 3).merge(t1.cell(0, 5)) # Merge 'จำแนกตามหมวดรายจ่าย' across 3 columns
    
    # Data Row
    t1.cell(2, 0).text = '1'
    t1.cell(2, 1).text = 'วัสดุปรับปรุงอาคารและสิ่งแวดล้อม'
    t1.cell(2, 2).text = '10,000'
    t1.cell(2, 3).text = '-'
    t1.cell(2, 4).text = '-'
    t1.cell(2, 5).text = '10,000'
    
    # Total Row
    t1.cell(3, 0).text = 'รวม'
    t1.cell(3, 1).text = 'รวม'
    t1.cell(3, 2).text = '10,000'
    t1.cell(3, 3).text = '-'
    t1.cell(3, 4).text = '-'
    t1.cell(3, 5).text = '10,000'
    
    # Center all header cells
    for r in range(2):
        for c in range(6):
            t1.cell(r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if t1.cell(r, c).paragraphs[0].runs:
                t1.cell(r, c).paragraphs[0].runs[0].bold = True

    # 6. Table 2: Evaluation
    add_p('\n6.  การประเมินผล', bold=True)
    t2 = doc.add_table(rows=2, cols=4)
    t2.style = 'Table Grid'
    headers2 = ['ที่', 'ดัชนีบ่งชี้ความสำเร็จ', 'วิธีการ/ประเมินผล', 'เครื่องมือที่ใช้วัด/ประเมินผล']
    for i, h in enumerate(headers2):
        t2.cell(0, i).text = h
        t2.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t2.cell(0, i).paragraphs[0].runs[0].bold = True
    
    t2.cell(1, 0).text = '1'
    t2.cell(1, 1).text = 'อาคารสถานที่และสิ่งแวดล้อมมีความปลอดภัยและพร้อมใช้งาน ร้อยละ 95'
    t2.cell(1, 2).text = '1. ตรวจสอบสภาพอาคาร\n2. สังเกตการใช้งาน'
    t2.cell(1, 3).text = '1. แบบประเมินคุณภาพสถานศึกษา'

    # 7. Expected Results
    add_p('\n7.    ผลที่คาดว่าจะได้รับ', bold=True)
    add_p('              7.1 โรงเรียนมีสภาพแวดล้อมที่สวยงาม ปลอดภัย และเอื้อต่อการเรียนรู้ของนักเรียน')
    add_p('              7.2 อาคารสถานที่และห้องเรียนมีความมั่นคง แข็งแรง และพร้อมใช้งานตลอดเวลา')
    add_p('              7.3 ครูและนักเรียนมีความสุขและมีคุณภาพชีวิตที่ดีในสถานศึกษา')

    # 8. Signatures
    add_p('\n\n')
    add_p('                     ผู้เสนอโครงการ                                                ผู้เห็นชอบโครงการ      ')
    add_p('\n')
    add_p('                ( นายอุดร  มะโนเนือง)                                        (นายสมศักดิ์  อิปิน)       ')
    add_p('                 ตำแหน่ง ครูชำนาญการ                         ประธานคณะกรรมการสถานศึกษาขั้นพื้นฐาน')
    add_p('\n')
    add_p('           ผู้อนุมัติโครงการ', bold=True, align='center')
    add_p('\n')
    add_p('                                               (นายสงวน  จันทอน)', align='center')
    add_p('                             ผู้อำนวยการโรงเรียนบ้านบ้านแม่ทราย(คุรุราษฎร์เจริญวิทย์)', align='center')

    # Force font for all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'TH SarabunPSK'
                        run.font.size = Pt(16)

    doc.save(output_path)
    print(f'Document created: {output_path}')

if __name__ == "__main__":
    p = r'C:\Users\User\Desktop\New folder (3)\สรุปโครงการ_อาคารสถานที่_เป๊ะ_100.docx'
    create_manual_doc(p)
