import sys
import os
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_final_masterpiece_v3(output_path):
    doc = Document()
    
    # 1. Page Setup (Margins 1 inch)
    section = doc.sections[0]
    section.top_margin = Twips(1440)
    section.bottom_margin = Twips(1440)
    section.left_margin = Twips(1440)
    section.right_margin = Twips(1440)

    # 2. Global Font & Spacing Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH SarabunPSK'
    font.size = Pt(16)
    
    # Tighten spacing
    pf = style.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    rPr = style._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'TH SarabunPSK')
    rFonts.set(qn('w:hAnsi'), 'TH SarabunPSK')
    rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    rFonts.set(qn('w:cs'), 'TH SarabunPSK')
    rPr.append(rFonts)

    def add_p(text, bold=False, align=None, indent=False):
        p = doc.add_paragraph()
        if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if indent: p.paragraph_format.first_line_indent = Twips(720) # 0.5 inch indent
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        
        run = p.add_run(text)
        run.bold = bold
        return p

    # Header
    add_p('โครงการ                 การพัฒนาอาคารสถานที่และสิ่งแวดล้อมอย่างมีคุณภาพ', bold=True)
    add_p('แผนงาน                  งานบริหารทั่วไป', bold=True)
    add_p('สนองกลยุทธ์สถานศึกษา           กลยุทธ์ที่ 4', bold=True)
    add_p('มาตรฐานดำเนินการ             มฐ.ที่ 2.5', bold=True)
    add_p('ลักษณะโครงการ                โครงการต่อเนื่อง', bold=True)
    add_p('งบประมาณ                         10,000  บาท', bold=True)
    add_p('ผู้รับผิดชอบโครงการ          นายอุดร  มะโนเนือง', bold=True)
    add_p('ระยะเวลาดำเนินการ            ปีการศึกษา 2569', bold=True)
    add_p('********************************************', align='center')

    # 1. Expanded Principles & Reasons
    add_p('1.  หลักการและเหตุผล', bold=True)
    p1_text = (
        'อาคารสถานที่และสิ่งแวดล้อมในสถานศึกษาเป็นปัจจัยพื้นฐานที่สำคัญอย่างยิ่งในการจัดการศึกษาที่มีคุณภาพ '
        'เนื่องจากสภาพแวดล้อมที่สะอาด ร่มรื่น มั่นคง แข็งแรง และปลอดภัย จะช่วยสร้างบรรยากาศที่เอื้อต่อการเรียนรู้ของนักเรียน '
        'และส่งเสริมประสิทธิภาพในการทำงานของคณะครูและบุคลากรทางการศึกษา ในแต่ละปีการศึกษา สถานศึกษาจะได้รับการสนับสนุนงบประมาณ '
        'เพื่อให้บริหารจัดการกลุ่มงานทั้ง 4 ด้าน ได้แก่ งานวิชาการ งานงบประมาณ งานบริหารงานบุคคล และงานบริหารทั่วไป '
        'ซึ่งงานพัฒนาอาคารสถานที่และสิ่งแวดล้อมถือเป็นภารกิจหลักในส่วนงานบริหารทั่วไปที่มีผลกระทบโดยตรงต่อความปลอดภัยและสุขภาวะของนักเรียน\n'
        '             โรงเรียนบ้านแม่ทรายจึงเล็งเห็นความจำเป็นในการยกระดับสภาพแวดล้อมและอาคารสถานที่ ห้องเรียน รวมถึงห้องน้ำสถานศึกษา '
        'ให้มีความพร้อมใช้งาน มีความเป็นระเบียบเรียบร้อย และสอดคล้องกับมาตรฐานการจัดการศึกษาที่กำหนดไว้ '
        'โดยเน้นการมีส่วนร่วมของคณะครู คณะกรรมการสถานศึกษา และชุมชน เพื่อให้โรงเรียนเป็นสถานที่ที่ปลอดภัย น่าอยู่ น่าเรียน '
        'และเป็นแหล่งเรียนรู้ที่มีคุณภาพสำหรับเยาวชนในพื้นที่อย่างแท้จริง จึงได้จัดทำโครงการนี้ขึ้น'
    )
    add_p(p1_text, indent=True)

    # 2. Objectives
    add_p('2.  วัตถุประสงค์', bold=True)
    add_p('            2.1  เพื่อปรับปรุงและพัฒนาอาคารสถานที่และสิ่งแวดล้อมให้มีความมั่นคง แข็งแรง และปลอดภัย', indent=False)
    add_p('            2.2  เพื่อสร้างบรรยากาศและสิ่งแวดล้อมที่เอื้อต่อการจัดการเรียนรู้ของนักเรียน', indent=False)
    add_p('            2.3  เพื่อให้สถานศึกษามีความพร้อมในการให้บริการแก่ชุมชนและหน่วยงานภายนอก', indent=False)

    # 3. Goals
    add_p('3.   เป้าหมาย', bold=True)
    add_p('     3.1  เชิงปริมาณ', bold=True)
    add_p('                   3.1.1  อาคารสถานที่และห้องเรียนได้รับการปรับปรุงและซ่อมแซมให้พร้อมใช้งาน ร้อยละ 95')
    add_p('                   3.1.2  สภาพแวดล้อมโดยรอบมีความสะอาด ร่มรื่น และสวยงาม ร้อยละ 95')
    add_p('     3.2  เชิงคุณภาพ', bold=True)
    add_p('                    3.2.1 โรงเรียนมีสภาพแวดล้อมที่เอื้อต่อการเรียนรู้ มีความปลอดภัย และเป็นที่พึงพอใจของผู้รับบริการ')

    # 4. Table 0: Activities (Exact Content)
    add_p('4. กิจกรรมและขั้นตอนการดำเนินโครงการ', bold=True)
    t0 = doc.add_table(rows=2, cols=4)
    t0.style = 'Table Grid'
    headers0 = ['ขั้นตอนการดำเนินงาน', 'ระยะเวลา', 'ผู้รับผิดชอบ', 'การประเมินผล']
    for i, h in enumerate(headers0):
        t0.cell(0, i).text = h
        t0.cell(0, i).paragraphs[0].runs[0].bold = True
        t0.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    step_content = (
        '1. ขั้นวางแผน(P)\n'
        '1.1 ประชุมคณะกรรมการและผู้เกี่ยวข้องเพื่อหาแนวทางในการจัดทำโครงการและวางแผนการดำเนินงาน\n'
        '1.2 แต่งตั้งคณะทำงาน\n'
        '1.3 เขียนโครงการ/กำหนดวันเวลาและนำเสนอโครงการเพื่อขออนุมัติ\n'
        '1.4 ประชุมคณะทำงานวางแผนการดำเนินงาน\n'
        '1.5 จัดทำแบบสอบถามความพึงพอใจผู้มีส่วนเกี่ยวข้อง\n'
        '2. ขั้นดำเนินการ(D)\n'
        ' ดำเนินงานตามโครงการ\n'
        '2.1 การจัดทำและปรับปรุงอาคารสถานที่ ห้องเรียน และสิ่งแวดล้อมให้เป็นระบบ สามารถตรวจสอบได้\n'
        '3. ขั้นประเมินผล (C)\n'
        '3.1 นิเทศติดตามผลการดำเนินงาน\n'
        '3.2 สรุป รายงานผลการจัดกิจกรรมตามแผน\n'
        '4. ขั้นปรับปรุง/พัฒนา (A)\n'
        '4.1 ปรับปรุงและวิธีการจัดกิจกรรมที่มีผลการพัฒนาไม่บรรลุตามเป้าหมาย'
    )
    time_content = '1-30 ก.ค.\n2569\n\n\nตลอด\nปีการศึกษา\n\n\n1-30 เม.ย.2570'
    owner_content = 'ผู้บริหาร\nนายอุดร มะโนเนือง'
    eval_content = 'วิธีการประเมินผล\n1. กรอกเอกสาร/สังเกต/สัมภาษณ์\n2. สรุปรายงานการพัฒนาอาคารสถานที่\n3. ประเมินภายใน/ภายนอก\n\nเครื่องมือประเมินผล\n1. สมุดบันทึกการนิเทศ\n2. บันทึกการนิเทศ\n3. แบบประเมินคุณภาพ'
    
    t0.cell(1, 0).text = step_content
    t0.cell(1, 1).text = time_content
    t0.cell(1, 2).text = owner_content
    t0.cell(1, 3).text = eval_content

    # 5. Table 1: Budget
    add_p('\n5.  งบประมาณ        งบอุดหนุนรายหัว 10,000 บาท', bold=True)
    t1 = doc.add_table(rows=4, cols=6)
    t1.style = 'Table Grid'
    t1.cell(0, 0).text = 'ที่'; t1.cell(1, 0).text = 'ที่'
    t1.cell(0, 1).text = 'รายการ'; t1.cell(1, 1).text = 'รายการ'
    t1.cell(0, 2).text = 'งบประมาณ'; t1.cell(1, 2).text = 'งบประมาณ'
    t1.cell(0, 3).text = 'จำแนกตามหมวดรายจ่าย'; t1.cell(1, 3).text = 'ตอบแทน'
    t1.cell(1, 4).text = 'ใช้สอย'; t1.cell(1, 5).text = 'วัสดุ'
    t1.cell(0, 0).merge(t1.cell(1, 0)); t1.cell(0, 1).merge(t1.cell(1, 1)); t1.cell(0, 2).merge(t1.cell(1, 2)); t1.cell(0, 3).merge(t1.cell(0, 5))
    t1.cell(2, 0).text = '1'; t1.cell(2, 1).text = 'วัสดุปรับปรุงอาคารและสิ่งแวดล้อม'; t1.cell(2, 2).text = '10,000'; t1.cell(2, 3).text = '-'; t1.cell(2, 4).text = '-'; t1.cell(2, 5).text = '10,000'
    t1.cell(3, 0).text = 'รวม'; t1.cell(3, 1).text = 'รวม'; t1.cell(3, 2).text = '10,000'; t1.cell(3, 3).text = '-'; t1.cell(3, 4).text = '-'; t1.cell(3, 5).text = '10,000'
    
    # Signatures
    add_p('\n\n')
    add_p('                     ผู้เสนอโครงการ                                                ผู้เห็นชอบโครงการ      ')
    add_p('\n')
    add_p('                ( นายอุดร  มะโนเนือง)                                        (นายสมศักดิ์  อิปิน)       ')
    add_p('                 ตำแหน่ง ครูชำนาญการ                         ประธานคณะกรรมการสถานศึกษาขั้นพื้นฐาน')
    add_p('\n')
    add_p('           ผู้อนุมัติโครงการ', bold=True, align='center')
    add_p('\n')
    add_p('                                               (นายสงวน  จันทอน)', align='center')
    add_p('                             ผู้อำนวยการโรงเรียนบ้านบ้านแม่ทราย(คุรุราษฎร์เจริญวิทย์)', align='center')

    # Global Font & Spacing Enforcement
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = 'TH SarabunPSK'
                        run.font.size = Pt(16)

    doc.save(output_path)
    print(f'Final Masterpiece v3 created: {output_path}')

if __name__ == "__main__":
    p = r'C:\Users\User\Desktop\New folder (3)\สรุปโครงการ_อาคารสถานที่_เป๊ะที่สุด_v3.docx'
    create_final_masterpiece_v3(p)
