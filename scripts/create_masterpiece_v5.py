import os
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_final_masterpiece_v5(output_path):
    doc = Document()
    
    # 1. Page Setup (914400 EMUs = 1440 Twips = 1 inch)
    section = doc.sections[0]
    section.top_margin = Twips(1440)
    section.bottom_margin = Twips(1440)
    section.left_margin = Twips(1440)
    section.right_margin = Twips(1440)

    # 2. Global Font & Spacing Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH SarabunPSK'
    font.size = Pt(16)
    
    pf = style.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.widow_control = False # Essential for Thai wrapping flow

    # Force Thai Font in XML
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = style._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'TH SarabunPSK')
    rFonts.set(qn('w:hAnsi'), 'TH SarabunPSK')
    rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    rFonts.set(qn('w:cs'), 'TH SarabunPSK')
    rPr.append(rFonts)

    def add_p(text, bold=False, align=None, indent=False):
        p = doc.add_paragraph()
        if align == 'center': 
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'left':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Use THAI_JUSTIFY if available, else JUSTIFY
            p.alignment = getattr(WD_ALIGN_PARAGRAPH, 'THAI_JUSTIFY', WD_ALIGN_PARAGRAPH.JUSTIFY)
        
        if indent: 
            p.paragraph_format.first_line_indent = Twips(720) # 0.5 inch indent
        
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        
        run = p.add_run(text)
        run.bold = bold
        return p

    # --- CONTENT START ---
    # Header Block
    add_p('โครงการ                 การพัฒนาอาคารสถานที่และสิ่งแวดล้อมอย่างมีคุณภาพ', bold=True, align='left')
    add_p('แผนงาน                  งานบริหารทั่วไป', bold=True, align='left')
    add_p('สนองกลยุทธ์สถานศึกษา           กลยุทธ์ที่ 4', bold=True, align='left')
    add_p('มาตรฐานดำเนินการ             มฐ.ที่ 2.5', bold=True, align='left')
    add_p('ลักษณะโครงการ                โครงการต่อเนื่อง', bold=True, align='left')
    add_p('งบประมาณ                         10,000  บาท', bold=True, align='left')
    add_p('ผู้รับผิดชอบโครงการ          นายอุดร  มะโนเนือง', bold=True, align='left')
    add_p('ระยะเวลาดำเนินการ            ปีการศึกษา 2569', bold=True, align='left')
    add_p('********************************************', align='center')

    # 1. Principles & Reasons (USER SPECIFIED FULL TEXT)
    add_p('1.  หลักการและเหตุผล', bold=True, align='left')
    p1_text = (
        'อาคารสถานที่และสิ่งแวดล้อมในสถานศึกษาเป็นปัจจัยพื้นฐานที่สำคัญอย่างยิ่งในการจัดการศึกษาที่มีคุณภาพ'
        'เนื่องจากสภาพแวดล้อมที่สะอาดร่มรื่นมั่นคงแข็งแรงและปลอดภัยจะช่วยสร้างบรรยากาศที่เอื้อต่อการเรียนรู้ของนักเรียน'
        'และส่งเสริมประสิทธิภาพในการทำงานของคณะครูและบุคลากรทางการศึกษา ในแต่ละปีการศึกษา สถานศึกษาจะได้รับการสนับสนุนงบประมาณ '
        'เพื่อให้บริหารจัดการกลุ่มงานทั้ง 4 ด้าน ได้แก่ งานวิชาการ งานงบประมาณงานบริหารงานบุคคลและงานบริหารทั่วไป'
        'ซึ่งงานพัฒนาอาคารสถานที่และสิ่งแวดล้อมถือเป็นภารกิจหลักในส่วนงานบริหารทั่วไปที่มีผลกระทบโดยตรงต่อความปลอดภัยและสุขภาวะของนักเรียน'
    )
    add_p(p1_text, indent=True)
    
    p2_text = (
        'โรงเรียนบ้านแม่ทรายจึงเล็งเห็นความจำเป็นในการยกระดับสภาพแวดล้อมและอาคารสถานที่ ห้องเรียนรวมถึงห้องน้ำสถานศึกษาให้มีความพร้อมใช้งาน'
        'มีความเป็นระเบียบเรียบร้อยและสอดคล้องกับมาตรฐานการจัดการศึกษาที่กำหนดไว้โดยเน้นการมีส่วนร่วมของคณะครู คณะกรรมการสถานศึกษา และชุมชน '
        'เพื่อให้โรงเรียนเป็นสถานที่ที่ปลอดภัยน่าอยู่น่าเรียนและเป็นแหล่งเรียนรู้ที่มีคุณภาพสำหรับเยาวชนในพื้นที่อย่างแท้จริง จึงได้จัดทำโครงการนี้ขึ้น'
    )
    add_p(p2_text, indent=True)

    # 2. Objectives
    add_p('2.  วัตถุประสงค์', bold=True, align='left')
    add_p('            2.1  เพื่อปรับปรุงและพัฒนาอาคารสถานที่และสิ่งแวดล้อมให้มีความมั่นคง แข็งแรง และปลอดภัย')
    add_p('            2.2  เพื่อสร้างบรรยากาศและสิ่งแวดล้อมที่เอื้อต่อการจัดการเรียนรู้ของนักเรียน')
    add_p('            2.3  เพื่อให้สถานศึกษามีความพร้อมในการให้บริการแก่ชุมชนและหน่วยงานภายนอก')

    # 3. Goals
    add_p('3.   เป้าหมาย', bold=True, align='left')
    add_p('     3.1  เชิงปริมาณ', bold=True, align='left')
    add_p('                   3.1.1  อาคารสถานที่และห้องเรียนได้รับการปรับปรุงและซ่อมแซมให้พร้อมใช้งาน ร้อยละ 95')
    add_p('                   3.1.2  สภาพแวดล้อมโดยรอบมีความสะอาด ร่มรื่น และสวยงาม ร้อยละ 95')
    add_p('     3.2  เชิงคุณภาพ', bold=True, align='left')
    add_p('                    3.2.1 โรงเรียนมีสภาพแวดล้อมที่เอื้อต่อการเรียนรู้ มีความปลอดภัย และเป็นที่พึงพอใจของผู้รับบริการ')

    # 4. Table 0: Activities (COMPLETE DETAILED CONTENT)
    add_p('4. กิจกรรมและขั้นตอนการดำเนินโครงการ', bold=True, align='left')
    t0 = doc.add_table(rows=2, cols=4)
    t0.style = 'Table Grid'
    headers0 = ['ขั้นตอนการดำเนินงาน', 'ระยะเวลา', 'ผู้รับผิดชอบ', 'การประเมินผล']
    for i, h in enumerate(headers0):
        t0.cell(0, i).text = h
        t0.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t0.cell(0, i).paragraphs[0].runs[0].bold = True
    
    step_content = (
        '1. ขั้นวางแผน(P)\n'
        '1.1 ประชุมคณะกรรมการและผู้เกี่ยวข้องเพื่อหาแนวทางในการจัดทำโครงการและวางแผนการดำเนินงาน\n'
        '1.2 แต่งตั้งคณะทำงาน\n'
        '1.3 เขียนโครงการ/กำหนดวันเวลาและนำเสนอโครงการเพื่อขออนุมัติ\n'
        '1.4 ประชุมคณะทำงานวางแผนการดำเนินงาน\n'
        '1.5 จัดทำแบบสอบถามความพึงพอใจผู้มีส่วนเกี่ยวข้อง\n'
        '2. ขั้นดำเนินการ(D)\n'
        ' ดำเนินงานตามโครงการ\n'
        '2.1 การจัดทำและปรับปรุงอาคารสถานที่ ห้องเรียน และสิ่งแวดล้อมให้เป็นระบบ สามารถตรวจสอบได้\n'
        '3. ขั้นประเมินผล (C)\n'
        '3.1 นิเทศติดตามผลการดำเนินงาน\n'
        '3.2 สรุป รายงานผลการจัดกิจกรรมตามแผน\n'
        '4. ขั้นปรับปรุง/พัฒนา (A)\n'
        '4.1 ปรับปรุงและวิธีการจัดกิจกรรมที่มีผลการพัฒนาไม่บรรลุตามเป้าหมาย'
    )
    t0.cell(1, 0).text = step_content
    t0.cell(1, 1).text = '1-30 ก.ค.\n2569\n\n\nตลอด\nปีการศึกษา\n\n\n1-30 เม.ย.2570'
    t0.cell(1, 2).text = 'ผู้บริหาร\nนายอุดร มะโนเนือง'
    t0.cell(1, 3).text = 'วิธีการประเมินผล\n1. กรอกเอกสาร/สังเกต/สัมภาษณ์\n2. จัดทำรายงานสรุป\n3. ประเมินภายใน/ภายนอก\n\nเครื่องมือประเมินผล\n1. สมุดบันทึกการนิเทศ\n2. บันทึกการนิเทศ\n3. แบบประเมินคุณภาพ'

    # 5. Table 1: Budget (Complex Headers)
    add_p('\n5.  งบประมาณ        งบอุดหนุนรายหัว 10,000 บาท', bold=True, align='left')
    t1 = doc.add_table(rows=4, cols=6)
    t1.style = 'Table Grid'
    t1.cell(0, 0).text = 'ที่'; t1.cell(1, 0).text = 'ที่'
    t1.cell(0, 1).text = 'รายการ'; t1.cell(1, 1).text = 'รายการ'
    t1.cell(0, 2).text = 'งบประมาณ'; t1.cell(1, 2).text = 'งบประมาณ'
    t1.cell(0, 3).text = 'จำแนกตามหมวดรายจ่าย'; t1.cell(1, 3).text = 'ตอบแทน'
    t1.cell(1, 4).text = 'ใช้สอย'; t1.cell(1, 5).text = 'วัสดุ'
    t1.cell(0, 0).merge(t1.cell(1, 0)); t1.cell(0, 1).merge(t1.cell(1, 1)); t1.cell(0, 2).merge(t1.cell(1, 2)); t1.cell(0, 3).merge(t1.cell(0, 5))
    t1.cell(2, 0).text = '1'; t1.cell(2, 1).text = 'วัสดุปรับปรุงอาคารและสิ่งแวดล้อม'; t1.cell(2, 2).text = '10,000'; t1.cell(2, 3).text = '-'; t1.cell(2, 4).text = '-'; t1.cell(2, 5).text = '10,000'
    t1.cell(3, 0).text = 'รวม'; t1.cell(3, 1).text = 'รวม'; t1.cell(3, 2).text = '10,000'; t1.cell(3, 3).text = '-'; t1.cell(3, 4).text = '-'; t1.cell(3, 5).text = '10,000'
    
    # 6. Table 2: Evaluation
    add_p('\n6.  การประเมินผล', bold=True, align='left')
    t2 = doc.add_table(rows=2, cols=4)
    t2.style = 'Table Grid'
    headers2 = ['ที่', 'ดัชนีบ่งชี้ความสำเร็จ', 'วิธีการ/ประเมินผล', 'เครื่องมือที่ใช้วัด/ประเมินผล']
    for i, h in enumerate(headers2):
        t2.cell(0, i).text = h
        t2.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        t2.cell(0, i).paragraphs[0].runs[0].bold = True
    t2.cell(1, 0).text = '1'; t2.cell(1, 1).text = 'ความปลอดภัยและพร้อมใช้งาน ร้อยละ 95'; t2.cell(1, 2).text = 'ตรวจสอบ/สังเกต'; t2.cell(1, 3).text = 'แบบประเมินคุณภาพ'

    # 7. Expected Results
    add_p('\n7.    ผลที่คาดว่าจะได้รับ', bold=True, align='left')
    add_p('              7.1 โรงเรียนมีสภาพแวดล้อมที่สวยงาม ปลอดภัย และเอื้อต่อการเรียนรู้ของนักเรียน')
    add_p('              7.2 อาคารสถานที่และห้องเรียนมีความมั่นคง แข็งแรง และพร้อมใช้งานตลอดเวลา')

    # 8. Signatures
    add_p('\n\n')
    add_p('                     ผู้เสนอโครงการ                                                ผู้เห็นชอบโครงการ      ', align='center')
    add_p('\n')
    add_p('                ( นายอุดร  มะโนเนือง)                                        (นายสมศักดิ์  อิปิน)       ', align='center')
    add_p('                 ตำแหน่ง ครูชำนาญการ                         ประธานคณะกรรมการสถานศึกษาขั้นพื้นฐาน', align='center')
    add_p('\n')
    add_p('           ผู้อนุมัติโครงการ', bold=True, align='center')
    add_p('\n')
    add_p('                                               (นายสงวน  จันทอน)', align='center')
    add_p('                             ผู้อำนวยการโรงเรียนบ้านบ้านแม่ทราย(คุรุราษฎร์เจริญวิทย์)', align='center')

    # Global Formatting Sync
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = 'TH SarabunPSK'
                        run.font.size = Pt(16)

    doc.save(output_path)
    print(f'Final Masterpiece v5 (FULL & PEUK) created: {output_path}')

if __name__ == "__main__":
    p = r'C:\Users\User\Desktop\New folder (3)\สรุปโครงการ_อาคารสถานที่_สมบูรณ์แบบ_v5.docx'
    create_final_masterpiece_v5(p)
