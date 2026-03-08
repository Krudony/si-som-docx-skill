import sys
import os
from docx import Document
from docx.shared import Twips, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class SiSomDocxEngine:
    def __init__(self, filename=None, font_size=16):
        if filename and os.path.exists(filename):
            self.doc = Document(filename)
        else:
            self.doc = Document()
            self._set_page_setup()
            self.set_font('TH SarabunPSK', font_size)

    def _set_page_setup(self):
        """Standard A4 with precise margins (1418 Twips)"""
        section = self.doc.sections[0]
        section.page_height = Twips(16838)
        section.page_width = Twips(11906)
        section.top_margin = Twips(1418)
        section.bottom_margin = Twips(1418)
        section.left_margin = Twips(1418)
        section.right_margin = Twips(1418)

    def set_font(self, name='TH SarabunPSK', size=16):
        """Set default font with Thai support."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = name
        font.size = Pt(size)
        rPr = style._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:eastAsia'), name)
        rFonts.set(qn('w:cs'), name)
        rPr.append(rFonts)

    def set_2_columns_continuous(self):
        """Add a continuous section break and set to 2 columns."""
        new_section = self.doc.add_section(WD_SECTION.CONTINUOUS)
        sectPr = new_section._sectPr
        cols = sectPr.xpath('./w:cols')
        if not cols:
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
        else:
            cols = cols[0]
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720') # 0.5 inch gap
        return new_section

    def add_header_center(self, text, bold=False):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bold and p.runs:
            p.runs[0].bold = True
        return p

    def save(self, filename):
        self.doc.save(filename)
        print(f"✅ Document saved: {filename}")

if __name__ == "__main__":
    print("Si-Som DOCX Engine v2.0 (Exam Ready)")
